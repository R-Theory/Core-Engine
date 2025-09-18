"""
Basic File Parsers - Common document parsing plugins

This module provides parsers for common file types including text, PDF, 
Word documents, and code files.
"""

import os
import mimetypes
from typing import List, Dict, Any
from pathlib import Path
import logging

from app.core.plugin_interface import (
    ParserPlugin, PluginResult, PluginMetadata, PluginType,
    PluginCapability, PluginConfig, PluginProcessingError
)

# Optional imports for different file types
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

logger = logging.getLogger(__name__)

class TextParserPlugin(ParserPlugin):
    """
    Parser for plain text files
    
    Supports: .txt, .md, .py, .js, .css, .html, .json, .xml, .csv, etc.
    """
    
    def __init__(self, config: PluginConfig):
        super().__init__(config)

    @property
    def metadata(self) -> PluginMetadata:
        return PluginMetadata(
            name="Text Parser",
            version="1.0.0",
            description="Parse plain text and code files",
            author="Core Engine Team",
            plugin_type=PluginType.PARSER,
            capabilities=[PluginCapability.PARSE, PluginCapability.READ],
            dependencies=[],
            config_schema={}
        )

    async def initialize(self) -> bool:
        """Initialize the text parser"""
        self.logger.info("Text parser initialized")
        return True

    async def process(self, data: Any, **kwargs) -> PluginResult:
        """Main processing method required by base Plugin class"""
        if isinstance(data, str):
            # Assume it's a file path
            return await self.parse(data)
        else:
            return PluginResult(
                success=False,
                error_message="TextParser expects a file path string"
            )

    async def can_parse(self, file_path: str, mime_type: str = None) -> bool:
        """Check if file can be parsed as text"""
        try:
            # Check by extension
            extension = Path(file_path).suffix.lower()
            if extension in self.get_supported_extensions():
                return True
            
            # Check by MIME type
            if mime_type and mime_type in self.get_supported_mime_types():
                return True
            
            # Try to detect text files by attempting to read them
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    # Try to read first 1KB to see if it's text
                    sample = f.read(1024)
                    # Check if it contains mostly printable characters
                    printable_ratio = sum(c.isprintable() or c.isspace() for c in sample) / len(sample)
                    return printable_ratio > 0.8
            except (UnicodeDecodeError, PermissionError):
                return False
                
        except Exception as e:
            self.logger.warning(f"Error checking if can parse {file_path}: {str(e)}")
            return False

    async def parse(self, file_path: str) -> PluginResult:
        """Parse text file and extract content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                        encoding_used = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return PluginResult(
                    success=False,
                    error_message="Unable to decode file with any supported encoding"
                )
            
            # Extract metadata
            file_path_obj = Path(file_path)
            title = file_path_obj.stem
            extension = file_path_obj.suffix.lower()
            
            # Detect language for code files
            language = self._detect_language(extension)
            
            # Count lines and characters
            lines = content.split('\n')
            line_count = len(lines)
            char_count = len(content)
            word_count = len(content.split())
            
            return PluginResult(
                success=True,
                data={
                    "title": title,
                    "content": content,
                    "file_type": "text",
                    "language": language,
                    "extension": extension
                },
                metadata={
                    "encoding": encoding_used,
                    "line_count": line_count,
                    "character_count": char_count,
                    "word_count": word_count,
                    "language": language,
                    "parser": "text"
                }
            )
            
        except Exception as e:
            self.logger.error(f"Failed to parse text file {file_path}: {str(e)}")
            return PluginResult(
                success=False,
                error_message=f"Text parsing error: {str(e)}"
            )

    def get_supported_extensions(self) -> List[str]:
        """Get supported file extensions"""
        return [
            '.txt', '.md', '.markdown', '.py', '.js', '.ts', '.jsx', '.tsx',
            '.css', '.scss', '.sass', '.html', '.htm', '.xml', '.json',
            '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf', '.log',
            '.csv', '.tsv', '.sql', '.sh', '.bash', '.ps1', '.bat',
            '.java', '.c', '.cpp', '.h', '.hpp', '.cs', '.php', '.rb',
            '.go', '.rs', '.swift', '.kt', '.scala', '.clj', '.hs',
            '.r', '.m', '.lua', '.pl', '.vim', '.dockerfile'
        ]

    def get_supported_mime_types(self) -> List[str]:
        """Get supported MIME types"""
        return [
            'text/plain', 'text/markdown', 'text/html', 'text/css',
            'text/javascript', 'text/xml', 'application/json',
            'application/yaml', 'text/csv', 'application/sql',
            'text/x-python', 'text/x-java-source', 'text/x-c',
            'text/x-c++', 'text/x-csharp', 'text/x-php', 'text/x-ruby',
            'text/x-go', 'text/x-rust', 'text/x-swift', 'text/x-kotlin'
        ]

    def _detect_language(self, extension: str) -> str:
        """Detect programming language from file extension"""
        language_map = {
            '.py': 'python',
            '.js': 'javascript',
            '.ts': 'typescript',
            '.jsx': 'jsx',
            '.tsx': 'tsx',
            '.java': 'java',
            '.c': 'c',
            '.cpp': 'cpp',
            '.h': 'c',
            '.hpp': 'cpp',
            '.cs': 'csharp',
            '.php': 'php',
            '.rb': 'ruby',
            '.go': 'go',
            '.rs': 'rust',
            '.swift': 'swift',
            '.kt': 'kotlin',
            '.scala': 'scala',
            '.html': 'html',
            '.css': 'css',
            '.scss': 'scss',
            '.sql': 'sql',
            '.sh': 'bash',
            '.bash': 'bash',
            '.md': 'markdown',
            '.json': 'json',
            '.xml': 'xml',
            '.yaml': 'yaml',
            '.yml': 'yaml'
        }
        return language_map.get(extension, 'text')

class PDFParserPlugin(ParserPlugin):
    """
    Parser for PDF files using pypdf
    """
    
    def __init__(self, config: PluginConfig):
        super().__init__(config)

    @property
    def metadata(self) -> PluginMetadata:
        return PluginMetadata(
            name="PDF Parser",
            version="1.0.0",
            description="Parse PDF documents and extract text content",
            author="Core Engine Team",
            plugin_type=PluginType.PARSER,
            capabilities=[PluginCapability.PARSE, PluginCapability.READ],
            dependencies=["pypdf"],
            config_schema={}
        )

    async def initialize(self) -> bool:
        """Initialize the PDF parser"""
        if not PDF_AVAILABLE:
            self.logger.error("pypdf not available")
            return False
        
        self.logger.info("PDF parser initialized")
        return True

    async def process(self, data: Any, **kwargs) -> PluginResult:
        """Main processing method required by base Plugin class"""
        if isinstance(data, str):
            # Assume it's a file path
            return await self.parse(data)
        else:
            return PluginResult(
                success=False,
                error_message="PDFParser expects a file path string"
            )

    async def can_parse(self, file_path: str, mime_type: str = None) -> bool:
        """Check if file is a PDF"""
        extension = Path(file_path).suffix.lower()
        return extension == '.pdf' or mime_type == 'application/pdf'

    async def parse(self, file_path: str) -> PluginResult:
        """Parse PDF file and extract text content"""
        try:
            if not PDF_AVAILABLE:
                return PluginResult(
                    success=False,
                    error_message="pypdf not available"
                )
            
            content_parts = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                doc_info = pdf_reader.metadata
                if doc_info:
                    metadata.update({
                        'pdf_title': doc_info.get('/Title', ''),
                        'pdf_author': doc_info.get('/Author', ''),
                        'pdf_subject': doc_info.get('/Subject', ''),
                        'pdf_creator': doc_info.get('/Creator', ''),
                        'pdf_producer': doc_info.get('/Producer', ''),
                        'pdf_creation_date': str(doc_info.get('/CreationDate', '')),
                        'pdf_modification_date': str(doc_info.get('/ModDate', ''))
                    })
                
                # Extract text from all pages
                page_count = len(pdf_reader.pages)
                metadata['page_count'] = page_count
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(f"--- Page {page_num} ---\n{page_text}")
                    except Exception as e:
                        self.logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                        continue
            
            # Combine all content
            content = '\n\n'.join(content_parts)
            
            # Use filename as title if no title in metadata
            title = metadata.get('pdf_title') or Path(file_path).stem
            
            # Calculate stats
            word_count = len(content.split()) if content else 0
            char_count = len(content)
            
            return PluginResult(
                success=True,
                data={
                    "title": title,
                    "content": content,
                    "file_type": "pdf"
                },
                metadata={
                    **metadata,
                    "word_count": word_count,
                    "character_count": char_count,
                    "parser": "pdf"
                }
            )
            
        except Exception as e:
            self.logger.error(f"Failed to parse PDF {file_path}: {str(e)}")
            return PluginResult(
                success=False,
                error_message=f"PDF parsing error: {str(e)}"
            )

    def get_supported_extensions(self) -> List[str]:
        """Get supported file extensions"""
        return ['.pdf']

    def get_supported_mime_types(self) -> List[str]:
        """Get supported MIME types"""
        return ['application/pdf']

class DocxParserPlugin(ParserPlugin):
    """
    Parser for Microsoft Word documents (.docx)
    """
    
    def __init__(self, config: PluginConfig):
        super().__init__(config)

    @property
    def metadata(self) -> PluginMetadata:
        return PluginMetadata(
            name="DOCX Parser",
            version="1.0.0",
            description="Parse Microsoft Word documents (.docx)",
            author="Core Engine Team",
            plugin_type=PluginType.PARSER,
            capabilities=[PluginCapability.PARSE, PluginCapability.READ],
            dependencies=["python-docx"],
            config_schema={}
        )

    async def initialize(self) -> bool:
        """Initialize the DOCX parser"""
        if not DOCX_AVAILABLE:
            self.logger.error("python-docx not available")
            return False
        
        self.logger.info("DOCX parser initialized")
        return True

    async def process(self, data: Any, **kwargs) -> PluginResult:
        """Main processing method required by base Plugin class"""
        if isinstance(data, str):
            # Assume it's a file path
            return await self.parse(data)
        else:
            return PluginResult(
                success=False,
                error_message="DocxParser expects a file path string"
            )

    async def can_parse(self, file_path: str, mime_type: str = None) -> bool:
        """Check if file is a DOCX document"""
        extension = Path(file_path).suffix.lower()
        return (extension == '.docx' or 
                mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    async def parse(self, file_path: str) -> PluginResult:
        """Parse DOCX file and extract text content"""
        try:
            if not DOCX_AVAILABLE:
                return PluginResult(
                    success=False,
                    error_message="python-docx not available"
                )
            
            # Load document
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content_parts.append('\t'.join(row_text))
            
            content = '\n'.join(content_parts)
            
            # Extract core properties if available
            metadata = {}
            try:
                core_props = doc.core_properties
                metadata.update({
                    'docx_title': core_props.title or '',
                    'docx_author': core_props.author or '',
                    'docx_subject': core_props.subject or '',
                    'docx_keywords': core_props.keywords or '',
                    'docx_comments': core_props.comments or '',
                    'docx_created': str(core_props.created) if core_props.created else '',
                    'docx_modified': str(core_props.modified) if core_props.modified else '',
                    'docx_last_modified_by': core_props.last_modified_by or ''
                })
            except Exception as e:
                self.logger.warning(f"Failed to extract document properties: {str(e)}")
            
            # Use document title or filename
            title = metadata.get('docx_title') or Path(file_path).stem
            
            # Calculate stats
            word_count = len(content.split()) if content else 0
            char_count = len(content)
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            return PluginResult(
                success=True,
                data={
                    "title": title,
                    "content": content,
                    "file_type": "docx"
                },
                metadata={
                    **metadata,
                    "word_count": word_count,
                    "character_count": char_count,
                    "paragraph_count": paragraph_count,
                    "table_count": len(doc.tables),
                    "parser": "docx"
                }
            )
            
        except Exception as e:
            self.logger.error(f"Failed to parse DOCX {file_path}: {str(e)}")
            return PluginResult(
                success=False,
                error_message=f"DOCX parsing error: {str(e)}"
            )

    def get_supported_extensions(self) -> List[str]:
        """Get supported file extensions"""
        return ['.docx']

    def get_supported_mime_types(self) -> List[str]:
        """Get supported MIME types"""
        return ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']